# -*- coding: utf-8 -*-
"""Build an editable Word (.docx) user manual with embedded real screenshots."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = r"d:/Vision-PSE/Contaminate_chechk_Size_web_ubuntu-main/Contaminate_chechk_Size_web_ubuntu-main/docs"
IMG = os.path.join(BASE, "manual_img")
OUT = os.path.join(BASE, "USER_MANUAL.docx")
FONT = "Leelawadee UI"   # Thai-capable font available on Windows

doc = Document()

def set_font(el_style):
    el_style.font.name = FONT
    el_style.font.size = Pt(11)
    rpr = el_style.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(a), FONT)

set_font(doc.styles['Normal'])
for s in ('Heading 1', 'Heading 2', 'Title'):
    try: set_font(doc.styles[s])
    except Exception: pass

def img(name, width=6.3):
    p = os.path.join(IMG, name)
    if os.path.exists(p):
        par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(p, width=Inches(width))
    else:
        doc.add_paragraph(f"[missing image: {name}]")

def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def p(t, bold=False):
    par = doc.add_paragraph(); r = par.add_run(t); r.bold = bold; return par
def bullet(t): doc.add_paragraph(t, style='List Bullet')
def num(t): doc.add_paragraph(t, style='List Number')

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    for i, hdr in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""; run = c.paragraphs[0].add_run(hdr); run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    doc.add_paragraph()

# ---------- Title ----------
title = doc.add_heading('PSE Vision — คู่มือการใช้งาน', level=0)
sub = doc.add_paragraph('Object Measurement System — ระบบวัดขนาดวัตถุ')
sub.runs[0].italic = True
note = doc.add_paragraph()
rn = note.add_run('✅ การตั้งค่าทั้งหมดในคู่มือนี้ทำได้จริงในโปรแกรม — บางส่วนใช้ได้ทันที (offline) '
                  'และบางส่วนต้องเปิดกล้องก่อน (ระบุไว้ในแต่ละหัวข้อ)')
rn.bold = True
doc.add_paragraph()

# ---------- 1 ----------
h1('1. ภาพรวมหน้าจอหลัก')
img('01_overview.png')
p('หน้าจอแบ่งเป็น 3 ส่วน:')
table(['ส่วน', 'รายละเอียด'], [
    ['ซ้าย — Sidebar', 'ปุ่มควบคุมกล้อง, นำเข้าภาพ, และสถานะระบบ'],
    ['กลาง — Live Camera Feed', 'ภาพสดจากกล้อง + กรอบ ROI ที่ตรวจจับ'],
    ['ขวา — Contour / System Log', 'ผลการตรวจจับ/วัดขนาด, สถิติ, และบันทึกระบบ'],
])
p('แถบด้านบน (Header):', bold=True)
bullet('CONNECTED / DISCONNECTED — สถานะการเชื่อมต่อ backend/กล้อง')
bullet('ปุ่มสลับธีม (ดวงอาทิตย์/พระจันทร์) — สลับโหมดสว่าง (Light) / มืด (Dark)')
bullet('Settings — เมนูตั้งค่าทั้งหมด')
bullet('นาฬิกา — เวลาปัจจุบัน')
doc.add_paragraph()
p('เมนู Settings รวมทุกการตั้งค่า (นำเมาส์ไปวางที่ปุ่ม Settings):', bold=True)
img('03_settings_menu.png')
p('General Settings, Calibration, Database Connection, Machines, LOT Management')

# ---------- 2 ----------
h1('2. การเปิดกล้องใช้งาน')
img('02_sidebar.png', width=2.4)
p('ที่ Sidebar → CAMERA CONTROL:', bold=True)
num('กด "Start Camera" เพื่อเริ่มกล้อง — ปุ่มจะเปลี่ยนเป็น "Stop Camera" และภาพสดจะขึ้นกลางจอ')
num('ตรวจสอบที่ STATUS: Connection (Connected = สำเร็จ), FPS (เฟรมต่อวินาที), Objects (จำนวนวัตถุ)')
num('กด "Stop Camera" อีกครั้งเพื่อหยุด')
p('ทดสอบโดยไม่มีกล้อง (Offline) — ที่ IMPORT IMAGE (OFFLINE):', bold=True)
bullet('Choose Image — เลือกไฟล์ภาพเพื่อทดสอบการตรวจจับ')
bullet('Calibration (mm) — ใส่ค่าคาลิเบรชันถ้าต้องการแปลงเป็นขนาดจริง')
bullet('Clear — ล้างภาพ')
p('หมายเหตุ: ถ้าขึ้น "No OAK camera detected" ให้ตรวจสายกล้อง USB/PoE หรือตั้ง IP กล้องในหน้า Settings')

# ---------- 3 ----------
h1('3. การตรวจจับ (Contour Detection) - เลือกเครื่องจักรผ่านหน้าเว็บ')
p('ทำได้บนหน้าเว็บทั้งหมด ไม่ต้องใช้โปรแกรม Desktop - ที่การ์ด Contour Detection ด้านขวา', bold=True)
img('10_detection_web.png')
num('กด "Start Camera" (Sidebar) ให้ภาพขึ้นก่อน')
num('ที่การ์ด Contour Detection -> ช่อง "Machine (Target)" -> เลือกเครื่องจักร '
    '(แต่ละเครื่องมีขนาดเป้าหมายในตัว พอเลือกก็เริ่มตรวจจับทันที เช่น "CUTT Line A (600-1200 mm2)")')
num('เลือก "— ไม่เลือก / หยุดตรวจจับ —" = หยุดตรวจจับ')
num('(เสริมที่ Sidebar) ปุ่ม "ยางดำ/ยางขาว" สลับชนิดวัสดุ, "Multi: ON" ตรวจหลายวัตถุ')
p('ผลลัพธ์ที่การ์ด Contour Detection:', bold=True)
bullet('เครื่องจักร / Target Area / Tolerance - เป้าหมายที่ใช้ตัดสิน')
bullet('สถิติ: Total, Pass, Near, Fail')
bullet('กรอบ ROI บนภาพสด: เขียว = ผ่าน, เหลือง = ใกล้เคียง, แดง = ไม่ผ่าน')
p('หมายเหตุ: เครื่องจักรต้องตั้ง Area Config (min/max mm2) ก่อน ที่ Settings -> Machines (ดูข้อ 6)')
p('ปุ่ม "ถ่าย + บันทึก DB":', bold=True)
bullet('อยู่ใต้ dropdown - ถ่ายภาพพร้อมกรอบ ROI แล้วบันทึกผลลงฐานข้อมูล')
bullet('ใช้ได้เมื่อเลือกเครื่องจักรแล้ว (กำลังตรวจจับ)')
bullet('ต้องเปิด Settings -> Database -> Enable Database Storage ก่อน ถึงจะบันทึกลง DB')

# ---------- 4 ----------
h1('4. การคาลิเบรตขนาด (Calibration)')
p('เปิดจาก Settings → Calibration')
img('04_calibration.png')
p('ระบบใช้การคาลิเบรตอัตโนมัติ (Auto) ด้วยวัตถุอ้างอิง 2 ชิ้น:', bold=True)
num('STEP 1 — วางวัตถุอ้างอิง 2 ชิ้นหน้ากล้อง: Object 1 = 50 × 50 mm, Object 2 = 20 × 20 mm '
    '(ให้เห็นชัดทั้งสองชิ้น)')
num('ทำตามขั้นตอนถัดไป เพื่อให้ระบบคำนวณค่า mm ต่อ pixel อัตโนมัติ')
num('เมื่อสำเร็จ ระบบบันทึกค่าคาลิเบรชัน และวัดขนาดวัตถุเป็นมิลลิเมตรได้ทันที')
p('หมายเหตุ: ควรคาลิเบรตใหม่ทุกครั้งที่ขยับกล้อง/เปลี่ยนระยะ เพื่อความแม่นยำ')

# ---------- 5 ----------
h1('5. การตั้งค่าฐานข้อมูล (Database)')
p('เปิดจาก Settings → Database Connection')
img('05_database.png')
p('ใช้เก็บผลการวัดลง SQL Server:', bold=True)
num('ดูสถานะการเชื่อมต่อ (Not Connected / Connected)')
num('ติ๊ก "Enable Database Storage" เพื่อเปิดการบันทึก — จะมีช่องกรอก Host, Username, Password, '
    'Database, Table และการตั้งค่ารอบการวัด (measurement session)')
num('บันทึกการตั้งค่า')
p('หมายเหตุ: ถ้าไม่เปิดใช้ ระบบยังวัดขนาดได้ปกติ เพียงแต่ไม่บันทึกลงฐานข้อมูล')

# ---------- 6 ----------
h1('6. การจัดการเครื่องจักร (Machines) — ตั้งเกณฑ์ Pass/Fail')
p('เปิดจาก Settings → Machines')
img('06_machines.png')
p('ใช้กำหนดเกณฑ์ขนาดที่ยอมรับของแต่ละเครื่องจักร (ใช้ตัดสิน Pass/Fail):', bold=True)
num('กรอกข้อมูล: รหัสเครื่องจักร*, ชื่อเครื่องจักร*, รายละเอียด, สถานที่ตั้ง, สถานะ (Active/Inactive)')
num('ตั้ง Area Configuration: พื้นที่ต่ำสุด (mm²), พื้นที่สูงสุด (mm²), ความคลาดเคลื่อนที่ยอมรับ (±mm²)')
num('บันทึก — เกณฑ์นี้จะถูกใช้ตัดสิน Pass / Near / Fail')

# ---------- 7 ----------
h1('7. การจัดการ LOT')
p('เปิดจาก Settings → LOT Management')
img('07_lots.png')
p('ใช้จัดกลุ่มการผลิต/ล็อตงาน:', bold=True)
num('กรอก รหัส LOT* (เช่น LOT-001), ชื่อ LOT* (เช่น Batch A-2024-03)')
num('เลือกประเภท (Type)* เช่น Rubber Type A')
num('ใส่รายละเอียดและสถานะ (Active/Inactive)')
num('กด "เพิ่ม LOT" — รายการจะแสดงในส่วน "รายการ LOT" ด้านล่าง')

# ---------- 8 ----------
h1('8. การตั้งค่าทั่วไป (General Settings)')
p('เปิดจาก Settings → General Settings')
img('08_general_settings.png')
bullet('Connected Cameras — แสดงกล้อง Luxonis ที่ต่ออยู่ (กด Refresh เพื่อค้นหาใหม่)')
bullet('Resolution — ความละเอียดกล้อง (เช่น 1920×1080)')
bullet('FPS Limit — จำกัดเฟรมต่อวินาที (เช่น 30)')
bullet('Show Depth Map — เปิด/ปิดการแสดงแผนที่ความลึก')
bullet('Measurement Settings — การตั้งค่าการวัด (เลื่อนลงดูเพิ่ม)')
bullet('ปุ่มล่าง: Reset to Defaults (คืนค่าเริ่มต้น), Cancel, Save Settings')

# ---------- 9 ----------
h1('9. โหมดสว่าง / มืด และการส่งออกข้อมูล')
p('กดปุ่มดวงอาทิตย์/พระจันทร์บนแถบด้านบนเพื่อสลับธีม — ระบบจะจดจำค่าที่เลือกไว้')
img('09_dark_overview.png')
p('กดปุ่ม "Export Data" (มุมขวาล่าง) เพื่อบันทึกผลการวัดและบันทึกระบบเป็นไฟล์')

# ---------- Summary ----------
h1('สรุปตารางการตั้งค่าทั้งหมด')
table(['เมนู', 'ทำอะไรได้', 'ต้องเปิดกล้อง?'], [
    ['Start Camera', 'เปิด/ปิดกล้อง', '-'],
    ['Start Contour + ยางดำ/ขาว + Multi', 'ตรวจจับ + ตีกรอบ ROI', 'ใช่'],
    ['Import Image', 'ทดสอบด้วยภาพไฟล์', 'ไม่'],
    ['Calibration', 'ตั้งค่า mm/pixel (auto 2 วัตถุ)', 'ใช่'],
    ['Database Connection', 'เก็บผลลง SQL Server', 'ไม่'],
    ['Machines', 'ตั้งเกณฑ์ขนาด Pass/Fail', 'ไม่'],
    ['LOT Management', 'จัดการล็อตงาน', 'ไม่'],
    ['General Settings', 'ความละเอียด/FPS/Depth ฯลฯ', 'ไม่'],
    ['Theme / Export', 'สลับธีม / ส่งออกข้อมูล', 'ไม่'],
])

doc.save(OUT)
print("saved:", OUT)
